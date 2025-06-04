from typing import Dict, Any, List
from pathlib import Path
from docx import Document
from docx.shared import Inches
from .base import BaseReporter


class WordReporter(BaseReporter):
    """Word报告生成器"""

    def generate(
        self, results: Dict[str, List[Dict[str, Any]]], output_path: str
    ) -> None:
        """生成Word报告

        Args:
            results: 评估结果数据，格式为 {文件标识: [结果列表]}
            output_path: 输出文件路径
        """
        doc = Document()
        doc.add_heading("ROB2 评估报告", 0)

        for study_id, study_results in results.items():
            # 为每个研究添加一个章节
            doc.add_heading(f"研究：{study_id}", level=1)

            # 添加研究结果表格
            table = doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            headers = table.rows[0].cells
            headers[0].text = "评估维度"
            headers[1].text = "偏倚风险"

            for entry in study_results:
                row = table.add_row().cells
                domain = entry.get("domain", "N/A")
                risk = (
                    entry.get("overall", {}).get("risk", "N/A")
                    if domain != "Overall risk of bias"
                    else entry.get("judgement", {}).get("overall", "N/A")
                )

                row[0].text = domain
                row[1].text = risk

            # 添加分隔
            doc.add_paragraph()

        doc.save(output_path)
        print(f"Word报告已生成: {output_path}")
